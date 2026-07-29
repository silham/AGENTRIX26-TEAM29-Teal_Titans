"""Owner: M6. Bytes -> plain text, for admin-uploaded knowledge-base documents.

Handles the formats Sri Lankan government sources actually arrive in: text PDFs
(circulars, instruction sheets), scanned PDFs, DOCX forms, and photographs of
office notices. Everything funnels to one ``extract()`` call returning text plus
the method used, so the admin UI can show how a document was read.

Dispatch is by file EXTENSION, not Content-Type: browsers routinely send
``application/octet-stream`` for .md and .docx, so the declared MIME is only a
cross-check.

Failures raise ``ExtractionError`` with a message written for the admin, not the
developer — it is stored on the document row and rendered verbatim in the admin
table, so it must say what to do next.
"""
from __future__ import annotations

import io
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path

logger = logging.getLogger(__name__)


class ExtractionError(Exception):
    """Extraction failed, with an actionable admin-facing message."""


@dataclass
class Extracted:
    text: str
    method: str  # text | pdf_text | pdf_ocr | pdf_embedded_image_ocr | docx | image_ocr
    page_count: int | None = None
    warnings: list[str] = field(default_factory=list)


# Extension -> kind. Also serves as the upload allowlist (app/documents/validation.py).
EXT_KIND: dict[str, str] = {
    ".txt": "text",
    ".md": "text",
    ".markdown": "text",
    ".pdf": "pdf",
    ".docx": "docx",
    ".png": "image",
    ".jpg": "image",
    ".jpeg": "image",
    ".webp": "image",
    ".tif": "image",
    ".tiff": "image",
    ".bmp": "image",
}

MIN_CHARS_PER_PAGE = 100  # below this, a PDF page is treated as scanned
MAX_OCR_PAGES = 30
MIN_USEFUL_CHARS = 50
OCR_DPI = 200  # 300 is ~2x slower for marginal gain on government documents


def sniff_kind(filename: str, mime: str | None = None) -> str:
    """Map a filename to an extraction kind, or raise with the allowed list."""
    ext = Path(filename or "").suffix.lower()
    kind = EXT_KIND.get(ext)
    if kind is None:
        allowed = ", ".join(sorted(EXT_KIND))
        if ext == ".doc":
            raise ExtractionError(
                "Legacy .doc files are not supported. Save the file as .docx or "
                "print it to PDF and upload that instead."
            )
        raise ExtractionError(
            f"Unsupported file type '{ext or filename}'. Allowed: {allowed}"
        )
    if mime and kind == "pdf" and "pdf" not in mime.lower():
        logger.info("Extension says PDF but Content-Type is %r; trusting the extension.", mime)
    return kind


# ---------------------------------------------------------------------------
# Per-format handlers
# ---------------------------------------------------------------------------


def _extract_text(data: bytes) -> tuple[str, list[str]]:
    warnings: list[str] = []
    try:
        # utf-8-sig strips a BOM if present; plain utf-8 would leave ﻿ in
        # the first chunk and pollute the embedding.
        return data.decode("utf-8-sig"), warnings
    except UnicodeDecodeError:
        warnings.append("File is not valid UTF-8; decoded as latin-1. Check for garbled characters.")
        return data.decode("latin-1", errors="replace"), warnings


def _extract_docx(data: bytes) -> str:
    try:
        import docx
    except ImportError as exc:
        raise ExtractionError("DOCX support requires python-docx (pip install -r requirements.txt).") from exc

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise ExtractionError(f"Could not open this .docx file: {exc}") from exc

    parts: list[str] = [p.text for p in document.paragraphs if p.text.strip()]
    # Government forms put requirement lists in tables. Skipping tables would
    # silently drop the most important content on the page.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _ocr_image_bytes(image_bytes: bytes) -> str:
    """OCR a single image via the existing Tesseract wrapper."""
    from app.documents import ocr

    try:
        return ocr.extract_text(image_bytes)
    except RuntimeError as exc:
        raise ExtractionError(str(exc)) from exc


def _ocr_pdf(data: bytes, reader, page_count: int) -> Extracted:
    """OCR a scanned PDF, degrading through progressively weaker backends."""
    if page_count > MAX_OCR_PAGES:
        raise ExtractionError(
            f"This PDF has {page_count} pages and appears to be scanned; OCR is capped at "
            f"{MAX_OCR_PAGES} pages. Split the document and upload the parts separately."
        )

    warnings: list[str] = []

    # 1) Render each page to an image with poppler. Best fidelity.
    try:
        from pdf2image import convert_from_bytes

        images = convert_from_bytes(data, dpi=OCR_DPI, fmt="png")
        pages: list[str] = []
        for image in images:
            buf = io.BytesIO()
            image.save(buf, format="PNG")
            pages.append(_ocr_image_bytes(buf.getvalue()))
        text = "\n\n".join(p for p in pages if p.strip())
        if len(text) >= MIN_USEFUL_CHARS:
            return Extracted(text, "pdf_ocr", page_count, warnings)
        warnings.append("Page rendering produced very little text; trying embedded images.")
    except ExtractionError:
        raise
    except Exception as exc:  # noqa: BLE001 — poppler missing is the common case
        logger.info("pdf2image unavailable or failed (%r); trying embedded images.", exc)
        warnings.append(
            "poppler is not installed, so page rendering was skipped and embedded images "
            "were OCR'd instead (lower fidelity). Install poppler for better results: "
            "brew install poppler / apt-get install poppler-utils."
        )

    # 2) No poppler needed: most scanners emit one full-page image per PDF page,
    #    and pypdf can hand those to us directly.
    embedded: list[str] = []
    for page in reader.pages:
        try:
            for image in page.images:
                embedded.append(_ocr_image_bytes(image.data))
        except ExtractionError:
            raise
        except Exception as exc:  # noqa: BLE001 — one bad image must not kill the page
            logger.info("Could not read an embedded image (%r); skipping it.", exc)
    text = "\n\n".join(p for p in embedded if p.strip())
    if len(text) >= MIN_USEFUL_CHARS:
        return Extracted(text, "pdf_embedded_image_ocr", page_count, warnings)

    raise ExtractionError(
        "This PDF appears to be scanned and no OCR backend could read it. Install Tesseract "
        "and poppler (brew install tesseract poppler / apt-get install tesseract-ocr "
        "poppler-utils), then use Reindex. If they are installed, the scan quality may be too "
        "low — try a higher-resolution scan."
    )


def _extract_pdf(data: bytes) -> Extracted:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ExtractionError("PDF support requires pypdf (pip install -r requirements.txt).") from exc

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:
        raise ExtractionError(f"Could not open this PDF: {exc}") from exc

    if reader.is_encrypted:
        # An empty user password is common for "protected" government PDFs.
        try:
            reader.decrypt("")
        except Exception as exc:
            raise ExtractionError(
                "This PDF is password-protected. Remove the password and upload it again."
            ) from exc

    page_count = len(reader.pages)
    if page_count == 0:
        raise ExtractionError("This PDF has no pages.")

    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception as exc:  # noqa: BLE001 — a malformed page shouldn't kill the file
            logger.info("Text extraction failed on one page (%r); continuing.", exc)
            pages.append("")
    text = "\n\n".join(pages).strip()

    if len(text) / page_count >= MIN_CHARS_PER_PAGE:
        return Extracted(text, "pdf_text", page_count, [])

    logger.info(
        "PDF yielded %d chars over %d pages (< %d/page); treating as scanned.",
        len(text), page_count, MIN_CHARS_PER_PAGE,
    )
    return _ocr_pdf(data, reader, page_count)


# ---------------------------------------------------------------------------
# Shared cleanup
# ---------------------------------------------------------------------------

_HYPHEN_BREAK = re.compile(r"(\w)-\n(\w)")
_BLANK_RUN = re.compile(r"\n{3,}")
_TRAILING_WS = re.compile(r"[ \t]+\n")


def clean_text(text: str) -> str:
    """Normalise whitespace and rejoin hyphenated line breaks.

    Cheap, and it materially improves chunk quality on OCR'd documents where a
    single requirement is routinely split across three lines.
    """
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _HYPHEN_BREAK.sub(r"\1\2", text)
    text = _TRAILING_WS.sub("\n", text)
    text = _BLANK_RUN.sub("\n\n", text)
    return text.strip()


def extract(data: bytes, filename: str, mime: str | None = None) -> Extracted:
    """Extract plain text from an uploaded document.

    Raises ExtractionError with an admin-facing message on any failure.
    """
    if not data:
        raise ExtractionError("The uploaded file is empty.")

    kind = sniff_kind(filename, mime)
    warnings: list[str] = []

    if kind == "text":
        raw, warnings = _extract_text(data)
        result = Extracted(raw, "text", None, warnings)
    elif kind == "docx":
        result = Extracted(_extract_docx(data), "docx", None, [])
    elif kind == "pdf":
        result = _extract_pdf(data)
    elif kind == "image":
        result = Extracted(_ocr_image_bytes(data), "image_ocr", 1, [])
    else:  # pragma: no cover — sniff_kind only returns the four kinds above
        raise ExtractionError(f"Unsupported file kind '{kind}'.")

    result.text = clean_text(result.text)
    if len(result.text) < MIN_USEFUL_CHARS:
        raise ExtractionError(
            f"Only {len(result.text)} characters of readable text were found — too little to "
            "index. If this is a scan, try a higher-resolution image; if it is a form, upload "
            "the instruction sheet that accompanies it."
        )
    return result
